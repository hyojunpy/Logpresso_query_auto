from pathlib import Path
from tempfile import TemporaryDirectory
import time
import unittest

from docx import Document

from app.services.indexer import DocumentIndex


def write_docx(path: Path, body: str) -> None:
    document = Document()
    document.add_heading("table", level=1)
    document.add_paragraph(body)
    document.save(path)


class DocumentIndexTest(unittest.TestCase):
    def test_ensure_current_indexes_once_until_document_changes(self):
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            doc_path = root / "manual.docx"
            db_path = root / "index.db"
            write_docx(doc_path, "table duration=24h sample_logs")

            index = DocumentIndex(db_path)
            self.assertFalse(index.status(doc_path)["indexed"])

            index.ensure_current(doc_path)
            first_status = index.status(doc_path)
            self.assertTrue(first_status["indexed"])
            self.assertFalse(first_status["stale"])
            first_mtime = first_status["last_indexed_mtime"]

            index.ensure_current(doc_path)
            second_status = index.status(doc_path)
            self.assertEqual(second_status["last_indexed_mtime"], first_mtime)
            self.assertFalse(second_status["stale"])

            time.sleep(0.02)
            write_docx(doc_path, "table duration=24h changed_logs")
            self.assertTrue(index.status(doc_path)["stale"])

            index.ensure_current(doc_path)
            refreshed_status = index.status(doc_path)
            self.assertFalse(refreshed_status["stale"])
            self.assertNotEqual(refreshed_status["last_indexed_mtime"], first_mtime)


if __name__ == "__main__":
    unittest.main()
