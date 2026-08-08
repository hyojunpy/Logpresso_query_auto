import importlib.util
import json
import logging
from pathlib import Path
from tempfile import TemporaryDirectory
import unittest
from unittest.mock import patch
from uuid import UUID

from docx import Document

from app.core.logging import JsonFormatter


def write_docx(path: Path, body: str) -> None:
    document = Document()
    document.add_heading("table", level=1)
    document.add_paragraph(body)
    document.save(path)


@unittest.skipIf(importlib.util.find_spec("fastapi") is None, "fastapi is not installed")
class ApiTest(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        from fastapi.testclient import TestClient
        from app.api.main import app

        cls.client = TestClient(app)

    def test_health(self):
        response = self.client.get("/api/v1/health")
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["status"], "ok")

    def test_api_responses_include_security_headers(self):
        response = self.client.get("/api/v1/health")
        self.assertEqual(response.headers["x-content-type-options"], "nosniff")
        self.assertEqual(response.headers["x-frame-options"], "DENY")
        self.assertEqual(response.headers["referrer-policy"], "no-referrer")
        self.assertEqual(
            response.headers["permissions-policy"],
            "camera=(), microphone=(), geolocation=()",
        )

    def test_api_response_includes_unique_request_id(self):
        first = self.client.get("/api/v1/health").headers["x-request-id"]
        second = self.client.get("/api/v1/health").headers["x-request-id"]
        self.assertEqual(UUID(first).hex, first)
        self.assertEqual(UUID(second).hex, second)
        self.assertNotEqual(first, second)

    def test_request_log_contains_metadata_but_not_query_or_body(self):
        secret_request = "sensitive natural language request"
        with patch("app.api.main.logger.info") as log:
            response = self.client.post(
                "/api/v1/query/generate?debug=private-query-value",
                json={"request": secret_request, "context": {}},
            )
        self.assertIn(response.status_code, (200, 422))
        metadata = log.call_args.kwargs["extra"]
        self.assertEqual(metadata["method"], "POST")
        self.assertEqual(metadata["path"], "/api/v1/query/generate")
        self.assertNotIn(secret_request, str(metadata))
        self.assertNotIn("private-query-value", str(metadata))

    def test_request_log_formatter_emits_structured_json(self):
        record = logging.LogRecord(
            name="logpresso.request",
            level=logging.INFO,
            pathname=__file__,
            lineno=1,
            msg="request_completed",
            args=(),
            exc_info=None,
        )
        record.request_id = "abc123"
        record.method = "GET"
        record.path = "/api/v1/health"
        record.status_code = 200
        record.duration_ms = 1.25
        payload = json.loads(JsonFormatter().format(record))
        self.assertEqual(payload["event"], "request_completed")
        self.assertEqual(payload["request_id"], "abc123")
        self.assertEqual(payload["status_code"], 200)
        self.assertNotIn("pathname", payload)

    def test_cors_allows_configured_streamlit_origin(self):
        response = self.client.options(
            "/api/v1/query/generate",
            headers={
                "Origin": "http://localhost:8501",
                "Access-Control-Request-Method": "POST",
                "Access-Control-Request-Headers": "Content-Type",
            },
        )
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.headers["access-control-allow-origin"], "http://localhost:8501")

    def test_cors_does_not_allow_unconfigured_origin(self):
        response = self.client.options(
            "/api/v1/query/generate",
            headers={
                "Origin": "https://example.invalid",
                "Access-Control-Request-Method": "POST",
            },
        )
        self.assertNotIn("access-control-allow-origin", response.headers)

    def test_generate_success(self):
        response = self.client.post(
            "/api/v1/query/generate",
            json={
                "request": "최근 24시간 동안 firewall_logs에서 출발지 IP별 차단 건수를 집계해서 많은 순으로 20개 보여줘",
                "context": {
                    "known_tables": ["firewall_logs"],
                    "known_fields": ["src_ip", "action", "_time"],
                },
            },
        )
        body = response.json()
        self.assertEqual(response.status_code, 200)
        self.assertEqual(body["status"], "generated")
        self.assertIn("table duration=24h firewall_logs", body["query"])
        self.assertTrue(body["validation"]["valid"])

    def test_generate_needs_clarification(self):
        response = self.client.post(
            "/api/v1/query/generate",
            json={"request": "에러 로그 보여줘", "context": {}},
        )
        body = response.json()
        self.assertEqual(response.status_code, 200)
        self.assertEqual(body["status"], "needs_clarification")
        self.assertIsNone(body["query"])
        self.assertTrue(body["questions"])

    def test_generate_rejects_oversized_natural_language_request(self):
        response = self.client.post(
            "/api/v1/query/generate",
            json={"request": "x" * 4001, "context": {}},
        )
        self.assertEqual(response.status_code, 422)
        self.assertEqual(response.json()["detail"][0]["type"], "string_too_long")

    def test_validate_rejects_oversized_query(self):
        response = self.client.post(
            "/api/v1/query/validate",
            json={"query": "x" * 20001},
        )
        self.assertEqual(response.status_code, 422)
        self.assertEqual(response.json()["detail"][0]["type"], "string_too_long")

    def test_command_search_rejects_oversized_query(self):
        response = self.client.get("/api/v1/commands/search", params={"q": "x" * 501})
        self.assertEqual(response.status_code, 422)

    def test_validate_invalid_query(self):
        response = self.client.post("/api/v1/query/validate", json={"query": "unknown firewall_logs"})
        body = response.json()
        self.assertEqual(response.status_code, 200)
        self.assertFalse(body["valid"])
        self.assertTrue([error for error in body["errors"] if error["code"] == "unknown_command"])

    def test_validate_pipeline_and_quote_errors(self):
        response = self.client.post(
            "/api/v1/query/validate",
            json={"query": 'table firewall_logs || search message == "ERROR'},
        )
        body = response.json()
        self.assertEqual(response.status_code, 200)
        self.assertFalse(body["valid"])
        error_codes = {error["code"] for error in body["errors"]}
        self.assertIn("empty_pipeline_segment", error_codes)
        self.assertIn("unbalanced_quotes", error_codes)

    def test_validate_fulltext_time_option_error(self):
        response = self.client.post(
            "/api/v1/query/validate",
            json={"query": 'fulltext duration=24h from=20260801 to=20260802 "1.2.3.4"'},
        )
        body = response.json()
        self.assertEqual(response.status_code, 200)
        self.assertFalse(body["valid"])
        self.assertTrue([error for error in body["errors"] if error["code"] == "exclusive_time_options"])

    def test_validate_ignores_syntax_inside_quoted_string(self):
        response = self.client.post(
            "/api/v1/query/validate",
            json={"query": 'table firewall_logs | search message == "duration=24h fake_func(value)"'},
        )
        body = response.json()
        self.assertEqual(response.status_code, 200)
        self.assertTrue(body["valid"], body["errors"])
        self.assertEqual(body["functions"], [])
        warning_codes = {warning["code"] for warning in body["warnings"]}
        self.assertNotIn("unknown_option", warning_codes)
        self.assertNotIn("unknown_function", warning_codes)

    def test_commands_search_and_detail(self):
        search_response = self.client.get("/api/v1/commands/search", params={"q": "fulltext"})
        self.assertEqual(search_response.status_code, 200)
        self.assertTrue(any(item["entry_name"] == "fulltext" for item in search_response.json()))

        detail_response = self.client.get("/api/v1/commands/fulltext")
        self.assertEqual(detail_response.status_code, 200)
        self.assertEqual(detail_response.json()["entry_name"], "fulltext")
        self.assertTrue(detail_response.json()["results"])

    def test_generate_unsupported_when_document_index_empty(self):
        from app.core.config import settings

        original_db_path = settings.db_path
        try:
            with TemporaryDirectory(ignore_cleanup_errors=True) as tmp:
                settings.db_path = Path(tmp) / "empty.db"
                response = self.client.post(
                    "/api/v1/query/generate",
                    json={
                        "request": "firewall_logs 보여줘",
                        "context": {"known_tables": ["firewall_logs"], "known_fields": ["_time"]},
                    },
                )
        finally:
            settings.db_path = original_db_path
        body = response.json()
        self.assertEqual(response.status_code, 200)
        self.assertEqual(body["status"], "unsupported")
        self.assertIsNone(body["query"])
        self.assertIn("reason", body["debug"])

    def test_reindex_document_not_found(self):
        from app.core.config import settings

        original_doc_path = settings.doc_path
        try:
            settings.doc_path = Path("docs") / "missing.docx"
            response = self.client.post("/api/v1/documents/reindex")
        finally:
            settings.doc_path = original_doc_path
        self.assertEqual(response.status_code, 404)
        self.assertEqual(response.json()["detail"]["code"], "document_not_found")

    def test_documents_status_and_reindex_success(self):
        from app.core.config import settings

        original_doc_path = settings.doc_path
        original_db_path = settings.db_path
        try:
            with TemporaryDirectory() as tmp:
                root = Path(tmp)
                settings.doc_path = root / "manual.docx"
                settings.db_path = root / "index.db"
                write_docx(settings.doc_path, "table duration=24h sample_logs")

                initial_status = self.client.get("/api/v1/documents/status").json()
                self.assertFalse(initial_status["indexed"])

                reindex_response = self.client.post("/api/v1/documents/reindex")
                self.assertEqual(reindex_response.status_code, 200)
                reindex_body = reindex_response.json()
                self.assertEqual(reindex_body["status"], "indexed")
                self.assertGreater(reindex_body["chunk_count"], 0)

                refreshed_status = self.client.get("/api/v1/documents/status").json()
                self.assertTrue(refreshed_status["indexed"])
                self.assertFalse(refreshed_status["stale"])
                self.assertEqual(refreshed_status["document"], "manual.docx")
        finally:
            settings.doc_path = original_doc_path
            settings.db_path = original_db_path


if __name__ == "__main__":
    unittest.main()
