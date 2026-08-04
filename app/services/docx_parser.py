from __future__ import annotations

import hashlib
import re
from pathlib import Path

from docx import Document

from app.models.document import DocumentChunk


COMMAND_HINTS = {
    "table", "search", "stats", "rollup", "timechart", "sort", "limit", "set", "setq",
    "stream", "fulltext", "logger", "fields", "eval", "rename", "parse", "explode",
    "evtx-file", "eml-file", "lnk-file",
}

FUNCTION_HINTS = {"count", "sum", "avg", "min", "max", "first", "last", "ago", "now", "str", "dateadd", "range", "iprange"}


class DocxParser:
    def parse(self, path: str | Path) -> list[DocumentChunk]:
        path = Path(path)
        document = Document(path)
        chunks: list[DocumentChunk] = []
        chapter: str | None = None
        section: str | None = None
        buffer: list[tuple[int, str]] = []

        def flush() -> None:
            nonlocal buffer
            if not buffer:
                return
            content = "\n".join(text for _, text in buffer).strip()
            entry_name, entry_type, content_type = self._classify(content)
            chunks.append(
                DocumentChunk(
                    document=path.name,
                    chapter=chapter,
                    section=section,
                    entry_type=entry_type,
                    entry_name=entry_name,
                    content_type=content_type,
                    content=content,
                    paragraph_start=buffer[0][0],
                    paragraph_end=buffer[-1][0],
                    ordinal=len(chunks),
                    content_hash=hashlib.sha256(content.encode("utf-8")).hexdigest(),
                    options=self._extract_options(content),
                    functions=self._extract_functions(content),
                )
            )
            buffer = []

        for index, paragraph in enumerate(document.paragraphs):
            text = " ".join(paragraph.text.split())
            if not text:
                continue
            style = (paragraph.style.name or "").lower()
            if "heading" in style or self._looks_like_heading(text):
                flush()
                if chapter is None or self._is_chapter(text):
                    chapter = text
                section = text
                continue
            buffer.append((index, text))
            if sum(len(item[1]) for item in buffer) > 1200:
                flush()
        flush()

        for table_index, table in enumerate(document.tables):
            rows = self._safe_table_rows(table)
            if not rows:
                continue
            content = "\n".join(" | ".join(cell for cell in row if cell) for row in rows)
            entry_name, entry_type, content_type = self._classify(content)
            chunks.append(
                DocumentChunk(
                    document=path.name,
                    chapter=chapter,
                    section=f"table {table_index + 1}",
                    entry_type=entry_type,
                    entry_name=entry_name,
                    content_type=content_type,
                    content=content[:5000],
                    paragraph_start=0,
                    paragraph_end=0,
                    ordinal=len(chunks),
                    content_hash=hashlib.sha256(content.encode("utf-8")).hexdigest(),
                    options=self._extract_options(content),
                    functions=self._extract_functions(content),
                )
            )
        return chunks

    def _safe_table_rows(self, table) -> list[list[str]]:
        rows: list[list[str]] = []
        try:
            iterable = table.rows
            for row in iterable:
                rows.append([" ".join(cell.text.split()) for cell in row.cells])
        except Exception:
            return rows
        return rows

    def _looks_like_heading(self, text: str) -> bool:
        return bool(re.match(r"^\d+(\.\d+)*\s+\S+", text)) or text in COMMAND_HINTS

    def _is_chapter(self, text: str) -> bool:
        return bool(re.match(r"^\d+\s+\S+", text)) or len(text) < 40

    def _classify(self, content: str) -> tuple[str | None, str, str]:
        lowered = content.lower()
        tokens = re.findall(r"[a-zA-Z_][a-zA-Z0-9_-]*", lowered)
        for token in tokens:
            if token in COMMAND_HINTS:
                if "syntax" in lowered or "문법" in content:
                    return token, "command", "syntax"
                if "example" in lowered or "예" in content:
                    return token, "command", "example"
                return token, "command", "description"
            if token in FUNCTION_HINTS:
                return token, "function", "description"
        return None, "text", "description"

    def _extract_options(self, content: str) -> list[str]:
        options = set(re.findall(r"\b([A-Za-z_][A-Za-z0-9_]*)\s*=", content))
        options.update(re.findall(r"\[([A-Za-z_][A-Za-z0-9_]*)=", content))
        return sorted(option.lower() for option in options)

    def _extract_functions(self, content: str) -> list[str]:
        functions = set()
        for name in re.findall(r"\b([A-Za-z_][A-Za-z0-9_]*)\s*\(", content):
            lowered = name.lower()
            if lowered in FUNCTION_HINTS:
                functions.add(lowered)
        return sorted(functions)
